# -*- coding: utf-8 -*-
"""
Генерация DOCX «Особенности и функционал ServiceDesk Plus ManageEngine».
Запуск: python generate_docx.py
Требуется: pip install -r requirements-docx.txt
"""
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def set_cell_shading(cell, rgb_hex):
    """Заливка ячейки цветом (rgb_hex без #)."""
    shd = OxmlElement('w:shd')
    shd.set(qn('w:fill'), rgb_hex)
    cell._tc.get_or_add_tcPr().append(shd)

def add_paragraph(doc, text, style='Normal', bold=False):
    p = doc.add_paragraph(text, style=style)
    if bold:
        for r in p.runs:
            r.bold = True
    return p

def run():
    doc = Document()
    sect = doc.sections[0]
    sect.top_margin = Cm(1.5)
    sect.bottom_margin = Cm(1.5)
    sect.left_margin = Cm(2)
    sect.right_margin = Cm(2)

    # Заголовок документа
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title.add_run('Особенности и функционал\nServiceDesk Plus ManageEngine')
    r.bold = True
    r.font.size = Pt(22)
    r.font.name = 'Times New Roman'
    title.space_after = Pt(24)

    # ========== 1. О продукте ==========
    doc.add_heading('1. О продукте', level=0)
    doc.add_paragraph(
        'ServiceDesk Plus — это платформа управления ИТ-услугами (ITSM) и корпоративного сервисного управления (ESM) '
        'от компании ManageEngine (подразделение Zoho Corporation). Продукт используется по всему миру: более 180 000 клиентов '
        'ManageEngine, присутствие в 190+ странах. ServiceDesk Plus поддерживает лучшие практики ITIL®, управление инцидентами, '
        'проблемами, изменениями, релизами, активами, каталог услуг, CMDB и проектами.'
    )
    doc.add_paragraph(
        'Продукт поставляется в трёх редакциях (Standard, Professional, Enterprise) и может разворачиваться как в облаке (SaaS), '
        'так и локально (on-premises). Это позволяет выбирать модель лицензирования и размещения под задачи организации.'
    )
    doc.add_paragraph('Основные требования, которым соответствует продукт:', style='List Bullet').runs[0].bold = True
    doc.add_paragraph(
        'Веб-интерфейс с локализацией (в т.ч. на русский), размещение On-Premise и в облаке, лицензии только для техников '
        '(количество пользователей-клиентов не ограничено).', style='List Bullet'
    )
    doc.add_paragraph(
        'Открытый REST API ServiceDesk Plus обеспечивает бесшовную интеграцию с ERP, CRM, системами мониторинга и другими '
        'внешними приложениями, позволяя использовать платформу как центральный оркестратор сквозных процессов.', style='List Bullet'
    )
    doc.add_paragraph(
        'Поддержка SSO и веб-сервисов, работа на мобильных устройствах и наличие мобильного приложения.', style='List Bullet'
    )
    doc.add_paragraph(
        'Редактирование процессов и объектов — в графическом интерфейсе (drag & drop); для глубокой кастомизации доступны '
        'пользовательские функции и скрипты (на базе поддерживаемых технологий).', style='List Bullet'
    )
    doc.add_paragraph()

    # ========== 2. Три редакции ==========
    doc.add_heading('2. Три редакции продукта', level=0)
    doc.add_paragraph('Официально предлагаются три редакции с чётким разграничением возможностей.')
    table = doc.add_table(rows=4, cols=4)
    table.style = 'Light Grid Accent 1'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    headers = ['Редакция', 'Основной фокус', 'Ключевые возможности', 'Для кого']
    for i, h in enumerate(headers):
        table.rows[0].cells[i].text = h
        for p in table.rows[0].cells[i].paragraphs:
            for r in p.runs:
                r.bold = True
    # Standard
    table.rows[1].cells[0].text = 'Standard'
    table.rows[1].cells[1].text = 'Сервис-деск (help desk)'
    table.rows[1].cells[2].text = 'Управление инцидентами, портал самообслуживания, база знаний, SLA, отчёты, многоплощадочная поддержка, автоназначение заявок, интеграция с AD/LDAP, API. Без управления активами.'
    table.rows[1].cells[3].text = 'Небольшие команды, базовый приём и обработка заявок.'
    for c in table.rows[1].cells:
        set_cell_shading(c, 'E0F2FE')
    # Professional
    table.rows[2].cells[0].text = 'Professional'
    table.rows[2].cells[1].text = 'Сервис-деск + активы'
    table.rows[2].cells[2].text = 'Всё из Standard + управление ИТ-активами, автообнаружение устройств, сканирование по агентам и по сети, закупки и договоры, отчёты по активам. CMDB — как дополнение.'
    table.rows[2].cells[3].text = 'Организации, которым нужен учёт активов и лицензий вместе с заявками.'
    for c in table.rows[2].cells:
        set_cell_shading(c, 'BAE6FD')
    # Enterprise
    table.rows[3].cells[0].text = 'Enterprise'
    table.rows[3].cells[1].text = 'Полный ITSM'
    table.rows[3].cells[2].text = 'Всё из Professional + управление проблемами, управление изменениями и релизами, каталог услуг с многоуровневыми согласованиями, управление ИТ-проектами, полноценная CMDB.'
    table.rows[3].cells[3].text = 'Крупные организации и предприятия с полным циклом ITIL и проектами.'
    for c in table.rows[3].cells:
        set_cell_shading(c, 'A5F3FC')
    doc.add_paragraph()
    doc.add_paragraph(
        'Во всех редакциях доступны: преобразование писем в заявки (email-to-ticket), интеграция с Active Directory/LDAP, '
        'поддержка нескольких порталов (ESM), SLA, база знаний, портал самообслуживания, отчёты и API. '
        'Аддон UEM Remote Access Plus (удалённый доступ к рабочим станциям) доступен в Professional и Enterprise. Ограничения по числу активов зависят от редакции и размера лицензии.'
    )
    doc.add_paragraph()

    # ========== Функционал (полный текст как на сайте) ==========
    doc.add_heading('Функционал', level=0)

    doc.add_heading('Управление инцидентами', level=1)
    doc.add_paragraph('Приём и маршрутизация:', style='List Bullet').runs[0].bold = True
    for s in [
        'Регистрация обращений по нескольким каналам: портал самообслуживания, почта, телефон, системы мониторинга (в т.ч. Zabbix и др.).',
        'Настройка правил обработки почтовых уведомлений: регистрация нового обращения, переписка в рамках открытой заявки, создание заявок от систем мониторинга.',
        'Связывание обращений с элементами и категориями CI в CMDB.',
        'Настраиваемый жизненный цикл заявки, автоматическое назначение техников по правилам и балансировка нагрузки.',
        'Классификация по категориям, влиянию и срочности.',
        'Оператор задач в заявке.',
    ]:
        doc.add_paragraph(s, style='List Bullet')
    doc.add_paragraph('Эскалация, SLA и удобство работы:', style='List Bullet').runs[0].bold = True
    for s in [
        'Автоматизация эскалации (функциональная и иерархическая).',
        'Согласования на портале и по электронной почте.',
        'SLA по времени реакции и решения с эскалациями и уведомлениями.',
        'Выделение заявок цветом в зависимости от оставшегося времени по SLA.',
        'Поиск и фильтрация списков заявок.',
        'Интеграция с базой знаний и рекомендация статей БЗ при регистрации заявки пользователем.',
    ]:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_heading('Управление активами (ITAM)', level=1)
    doc.add_paragraph('Обнаружение и учёт:', style='List Bullet').runs[0].bold = True
    for s in [
        'Обнаружение устройств по сети и по агентам (Windows, macOS, Linux).',
        'Учёт оборудования и ПО, управление лицензиями и соответствием.',
        'Связь активов с заявками и CMDB для анализа влияния сбоев и планирования изменений.',
    ]:
        doc.add_paragraph(s, style='List Bullet')
    doc.add_paragraph('Закупки и договоры:', style='List Bullet').runs[0].bold = True
    for s in [
        'Управление закупками и договорами.',
        'Напоминания об истечении гарантий и контрактов.',
    ]:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_heading('CMDB (база конфигураций)', level=1)
    doc.add_paragraph('Модель и визуализация:', style='List Bullet').runs[0].bold = True
    for s in [
        'Ведение CMDB с различными типами связей CI, атрибутами, классификаторами и статусами (жизненный цикл CI).',
        'Визуализация взаимосвязей конфигурационных единиц.',
        'Динамическое взаимодействие с деревом (развёртка/свёртка, переход на нужный уровень).',
    ]:
        doc.add_paragraph(s, style='List Bullet')
    doc.add_paragraph('Расчёт стоимости и использование:', style='List Bullet').runs[0].bold = True
    for s in [
        'Расчёт и просмотр TCO для конфигурационных единиц.',
        'Автоматизация расчёта стоимости оказания услуг в разрезе месяца и года.',
        'Использование CMDB при анализе инцидентов, проблем и изменений для принятия решений по влиянию на инфраструктуру.',
    ]:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_heading('Каталог услуг', level=1)
    doc.add_paragraph('Публикация и привязки:', style='List Bullet').runs[0].bold = True
    for s in [
        'Публикация ИТ- и бизнес-услуг на портале самообслуживания с настраиваемыми шаблонами запросов.',
        'Привязка каталога к CMDB.',
        'Привязка контрактов к услугам (расчёт стоимости предоставления услуг).',
    ]:
        doc.add_paragraph(s, style='List Bullet')
    doc.add_paragraph('Согласования и доступ:', style='List Bullet').runs[0].bold = True
    for s in [
        'Многоуровневые согласования, SLA и задачи по каждому типу услуги.',
        'Персонализированный доступ к каталогу по ролям и правам.',
        'Возможность создания обращения или инцидента по доступным пользователю услугам.',
    ]:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_heading('Управление изменениями и релизами', level=1)
    doc.add_paragraph('Изменения:', style='List Bullet').runs[0].bold = True
    for s in [
        'Типы изменений (стандартные, срочные, критичные) с отдельными рабочими процессами.',
        'Визуальный конструктор этапов, согласований и уведомлений.',
        'CAB, планы внедрения и отката.',
        'Календарь изменений и проверка конфликтов.',
    ]:
        doc.add_paragraph(s, style='List Bullet')
    doc.add_paragraph('Релизы:', style='List Bullet').runs[0].bold = True
    for s in [
        'Управление релизами с шаблонами и этапами.',
        'Связь релизов с изменениями и активами.',
    ]:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_heading('База знаний и портал самообслуживания', level=1)
    doc.add_paragraph('База знаний:', style='List Bullet').runs[0].bold = True
    for s in [
        'База знаний на портале самообслуживания с полнотекстовым поиском, категориями и процессом согласования статей.',
        'Рекомендация статей БЗ при регистрации заявки пользователем.',
        'Публикация решений на портале для снижения потока типовых заявок.',
    ]:
        doc.add_paragraph(s, style='List Bullet')
    doc.add_paragraph('Портал и интерфейс:', style='List Bullet').runs[0].bold = True
    for s in [
        'Брендирование и изменение внешнего вида портала.',
        'Персонализация в зависимости от прав пользователя (роли): доступность и отображение списков услуг.',
        'События и новости на портале.',
        'Виджеты, каталог услуг, отслеживание статуса запросов.',
        'Настройка интерфейса в зависимости от роли пользователя.',
        'Поддержка мобильных приложений.',
    ]:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_heading('Отчётность и дашборды', level=1)
    doc.add_paragraph('Отчёты:', style='List Bullet').runs[0].bold = True
    for s in [
        'Формирование отчётов в графическом интерфейсе (настройка мышью).',
        'Отчёты в рамках разделов (инциденты, активы, изменения, SLA, загрузка техников), настраиваемые в графическом интерфейсе.',
        'Возможность самостоятельного написания SQL-запросов для отчётов (on-premises).',
        'Выгрузка отчётов в xls, pdf, xml.',
        'Единый отчёт по мониторингу SLA всех услуг.',
    ]:
        doc.add_paragraph(s, style='List Bullet')
    doc.add_paragraph('Дашборды и аналитика:', style='List Bullet').runs[0].bold = True
    for s in [
        'Планировщик рассылки отчётов по почте.',
        'Интерактивные дашборды с виджетами.',
        'Интеграция с Zoho Analytics / Analytics Plus для расширенной аналитики.',
    ]:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_heading('Enterprise Service Management (ESM)', level=1)
    doc.add_paragraph('Суть ESM:', style='List Bullet').runs[0].bold = True
    for s in [
        'Enterprise Service Management — управление корпоративным сервисом в едином информационном пространстве предприятия.',
        'Одновременная автоматизация нескольких направлений: ИТ-поддержка, оборудование, HR, финансы и др.',
        'Для каждого отдела или организации создаётся отдельный экземпляр службы поддержки со своими шаблонами, пользователями и рабочими процессами.',
        'Поддержка методологии ITIL для ИТ-подразделений.',
    ]:
        doc.add_paragraph(s, style='List Bullet')
    doc.add_paragraph('Портал ESM и экземпляры:', style='List Bullet').runs[0].bold = True
    for s in [
        'Портал ESM — единая консоль, в которой пользователи работают с разными экземплярами служб, создают и обрабатывают заявки, управляют настройками.',
        'Экземпляры могут работать независимо друг от друга, использовать общие справочники или обмениваться данными внутри системы.',
        'Каталог ESM: пользователи, экземпляры, права доступа, авторизация.',
    ]:
        doc.add_paragraph(s, style='List Bullet')

    doc.add_heading('Интеграции', level=1)
    doc.add_paragraph('ManageEngine и мониторинг:', style='List Bullet').runs[0].bold = True
    for s in [
        'Нативные интеграции: OpManager, Endpoint Central (Desktop Central), Applications Manager, ADManager Plus, ADSelfService Plus, Password Manager Pro, Site24x7.',
        'Настройка интеграции с Zabbix и правил регистрации инцидентов от систем мониторинга.',
    ]:
        doc.add_paragraph(s, style='List Bullet')
    doc.add_paragraph('Бизнес-системы и API:', style='List Bullet').runs[0].bold = True
    for s in [
        'Интеграция с Jira (создание задач по проектам, получение задач из проектов, создание проектов).',
        'Бизнес-интеграции: Outlook, Office 365, Microsoft Teams, календарь O365.',
        'Поддержка WebServices, REST API. Полный доступ к данным и автоматизации через API.',
    ]:
        doc.add_paragraph(s, style='List Bullet')
    doc.add_paragraph()

    # ========== 3. Облако и on-premises ==========
    doc.add_heading('3. Облачное и локальное размещение', level=0)
    doc.add_paragraph('ServiceDesk Plus доступен в двух вариантах: облако (SaaS) и локальное развёртывание (on-premises).')
    doc.add_paragraph('Облако (Cloud): подписка, хостинг в дата-центрах ManageEngine в разных регионах (США, ЕС, Австралия, Индия и др.).', style='List Bullet')
    doc.add_paragraph(
        'Локальное развёртывание (On-Premises): установка на инфраструктуре заказчика; лицензия бессрочная или годовая; '
        'полный контроль над данными. Требования: Windows или Linux; СУБД по документации.', style='List Bullet'
    )
    doc.add_paragraph()

    # ========== 4. MSP ==========
    doc.add_heading('4. ServiceDesk Plus MSP', level=0)
    doc.add_paragraph(
        'ServiceDesk Plus MSP — отдельный продукт для провайдеров управляемых услуг (MSP). Единая платформа объединяет '
        'возможности ITSM (в т.ч. ITAM, CMDB) и элементы PSA: настройка и предоставление сервисов для нескольких заказчиков с учётом прибыльности.'
    )
    for s in [
        'Мультитенантность: отдельные настройки автоматизации, SLA, базы знаний, активы и отчёты для каждого клиента.',
        'Мультисайтовая поддержка: управление несколькими площадками заказчиков из одной консоли с разными часовыми поясами, правилами и SLA.',
        'Биллинг: предопределённые сервисные планы и рекуррентные платежи.',
        'Интеграция с RMM-инструментами ManageEngine (например, RMM Central): единая консоль для мониторинга и сервис-деска.',
        'Автоматизация: визуальные конструкторы рабочих процессов без кода.',
        'ИИ-агент Zia доступен как первая линия поддержки (функция не адаптирована для российского рынка).',
    ]:
        doc.add_paragraph(s, style='List Bullet')
    doc.add_paragraph(
        'MSP доступен в облаке и on-premises; редакции и цены отличаются от «обычного» ServiceDesk Plus. '
        'Подробности: manageengine.com/products/service-desk-msp'
    )
    doc.add_paragraph()

    # ========== 5. Преимущества ==========
    doc.add_heading('5. Преимущества и ключевые особенности ServiceDesk Plus', level=0)
    for s in [
        'Масштабируемость по зрелости: три редакции позволяют начать с help desk и поэтапно подключать активы, затем изменения, проблемы, каталог услуг, проекты и CMDB.',
        'Неограниченное количество конечных пользователей и согласующих: лицензия считается на техников; запросчики и согласующие не лицензируются отдельно.',
        'Готовая автоматизация: почтовые правила, бизнес-правила, триггеры, правила полей и форм (создание и редактирование произвольных полей формы в зависимости от выбранных значений), настраиваемый жизненный цикл заявки, планировщик задач; для сложных сценариев — пользовательские функции и скрипты.',
        'Персонализация и настройка: интерфейсы в зависимости от пользователя, роли и места работы; управление ролями и правами доступа (по ролям и услугам); управление данными (услуги, роли, пользователи, группы); расширяемость системы.',
        'Интеграции: нативные интеграции с продуктами ManageEngine (OpManager, Endpoint Central, ADManager Plus, ADSelfService Plus, Password Manager Pro, Site24x7 и др.), а также с Outlook, Office 365, Microsoft Teams, Jira, полный доступ — через REST API.',
        'ИИ и аналитика: умная категоризация заявок, предсказание приоритета, автоназначение техников, виртуальный агент Zia, интеграция с Zoho Analytics/Analytics Plus. (Функции помощника Zia пока не адаптированы для использования в России.)',
        'Признание рынка: ManageEngine фигурирует в отчётах Gartner (Magic Quadrant для ITSM) и Forrester (Total Economic Impact).',
        'Функциональное мобильное приложение для техников и пользователей (в том числе от российских разработчиков).',
    ]:
        doc.add_paragraph(s, style='List Bullet')
    doc.add_paragraph()

    # ========== Ограничения ==========
    doc.add_heading('Ограничения и особенности', level=0)
    doc.add_paragraph(
        'Часть продвинутых возможностей ИИ и интеграций доступна только в облачной версии или в определённых регионах. '
        'В редакции Standard нет управления активами и аддона UEM Remote Access Plus (удалённый доступ к рабочим станциям).'
    )
    doc.add_paragraph()

    # ========== Глоссарий ==========
    doc.add_heading('Глоссарий', level=0)
    doc.add_paragraph('Краткие пояснения терминов.')
    gloss = [
        ('ITSM', 'Управление ИТ-услугами — процессы и практики для проектирования, предоставления и поддержки ИТ-услуг.'),
        ('ESM', 'Корпоративное сервисное управление — распространение ITSM на HR, финансы, эксплуатацию и т.д.'),
        ('ITIL®', 'Набор лучших практик по управлению ИТ-услугами; ServiceDesk Plus поддерживает процессы ITIL.'),
        ('CMDB', 'База конфигураций — элементы CI и связи для учёта инфраструктуры и влияния изменений.'),
        ('SLA', 'Соглашение об уровне обслуживания — сроки реакции и решения заявок с эскалацией.'),
        ('UEM Remote Access Plus', 'Аддон удалённого доступа к рабочим станциям из консоли ServiceDesk Plus (Endpoint Central); доступен в Professional и Enterprise.'),
        ('SaaS', 'ПО как услуга — облачное развёртывание с подпиской.'),
        ('On-premises', 'Локальное развёртывание на инфраструктуре заказчика.'),
        ('MSP', 'Managed Service Provider — провайдер управляемых услуг; версия ServiceDesk Plus MSP.'),
    ]
    t2 = doc.add_table(rows=1 + len(gloss), cols=2)
    t2.style = 'Light List Accent 1'
    t2.rows[0].cells[0].text = 'Термин'
    t2.rows[0].cells[1].text = 'Пояснение'
    for r in t2.rows[0].cells:
        for p in r.paragraphs:
            for run in p.runs:
                run.bold = True
    for i, (term, desc) in enumerate(gloss):
        t2.rows[i + 1].cells[0].text = term
        t2.rows[i + 1].cells[1].text = desc
    doc.add_paragraph()

    # ========== Источники ==========
    doc.add_heading('Источники и полезные ссылки', level=0)
    for s in [
        'Официальная страница продукта: https://www.manageengine.com/products/service-desk/',
        'Сравнение редакций: https://www.manageengine.com/products/service-desk/sdp-editions.html',
        'On-Premises vs. Cloud: https://www.manageengine.com/products/service-desk/hosted-on-premise-vs-saas-cloud.html',
        'ServiceDesk Plus MSP: https://www.manageengine.com/products/service-desk-msp/',
        'Клиенты: https://www.manageengine.com/customers.html',
    ]:
        doc.add_paragraph(s, style='List Bullet')

    out_path = 'ServiceDesk-Plus-обзор.docx'
    doc.save(out_path)
    print('Создан файл: %s' % out_path)

if __name__ == '__main__':
    run()
